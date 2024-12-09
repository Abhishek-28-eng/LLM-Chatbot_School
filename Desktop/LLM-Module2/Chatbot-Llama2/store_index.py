import os
from dotenv import load_dotenv
import pandas as pd
from langchain_community.vectorstores import Pinecone as LangchainPinecone
from langchain_huggingface import HuggingFaceEmbeddings
from pinecone import Pinecone, ServerlessSpec, Index
from src.helper import download_hugging_face_embeddings, text_split
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader
#from ocr_handler import extract_text_with_ocr  # Assuming OCR is implemented
from docx import Document as DocxDocument

# Load environment variables
load_dotenv()

# Pinecone configuration
PINECONE_API_KEY = os.environ.get('PINECONE_API_KEY')
PINECONE_API_ENV = os.environ.get('PINECONE_API_ENV')

# Folder containing your files
DATA_FOLDER = "data/"

# def extract_text_from_file(file_path):
#     file_extension = os.path.splitext(file_path)[1].lower()

#     print(f"Processing file: {file_path} (Extension: {file_extension})")

#     try:
#         if file_extension == ".pdf":
#             print(f"Performing OCR on PDF: {file_path}...")
#             text = extract_text_from_pdf_with_ocr(file_path)
#             return text

#         elif file_extension == ".docx":
#             doc = DocxDocument(file_path)
#             text = "\n".join([para.text for para in doc.paragraphs])
#             return text

#         elif file_extension == ".txt":
#             with open(file_path, "r", encoding="utf-8") as f:
#                 text = f.read()
#             return text

#         elif file_extension == ".csv":
#             df = pd.read_csv(file_path)
#             text = df.to_string(index=False)
#             return text

#         elif file_extension in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
#             print(f"Performing OCR on {file_path}...")
#             text = extract_text_with_ocr(file_path)
#             print(f"OCR Extracted Text: {text[:100]}...")  # Print first 100 characters for debugging
#             return text

#         else:
#             return f"Unsupported file format: {file_extension}"

#     except Exception as e:
#         return f"Error processing {file_path}: {str(e)}"

# Function to extract text from different file formats
def extract_text_from_file(file_path):
    """
    Extract text from a file based on its format.
    Supports PDF, DOCX, TXT, CSV, and image files.
    """
    file_extension = os.path.splitext(file_path)[1].lower()

    try:
        if file_extension == ".pdf":
            loader = PyPDFLoader(file_path)
            document = loader.load()
            text = " ".join([page.page_content for page in document])
            return text

        elif file_extension == ".docx":
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text

        elif file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            return text

        elif file_extension == ".csv":
            df = pd.read_csv(file_path)
            text = df.to_string(index=False)
            return text

        elif file_extension in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
            text = extract_text_with_ocr(file_path)
            return text

        else:
            return f"Unsupported file format: {file_extension}"

    except Exception as e:
        return f"Error processing {file_path}: {str(e)}"


# Extract text from all files in the 'data' folder
all_text = []
for file_name in os.listdir(DATA_FOLDER):
    file_path = os.path.join(DATA_FOLDER, file_name)
    if os.path.isfile(file_path):
        print(f"Processing file: {file_name}")
        extracted_text = extract_text_from_file(file_path)
        all_text.append(extracted_text)

# Convert the extracted text into Document objects
documents = [Document(page_content=text) for text in all_text if text]

# Split the documents into chunks
text_chunks = text_split(documents)

# Load embeddings from Hugging Face
embeddings = download_hugging_face_embeddings()

# Initialize Pinecone client
pc = Pinecone(api_key=PINECONE_API_KEY, environment=PINECONE_API_ENV)

# Define index parameters
index_name = "llm-chatbot"
dimension = 384  # Ensure this matches the embedding model's output dimension
metric = "cosine"

# Check if the index exists, create it if not
if index_name not in pc.list_indexes().names():
    pc.create_index(
        name=index_name,
        dimension=dimension,
        metric=metric,
        spec=ServerlessSpec(
            cloud="aws",        # Specify cloud provider
            region="us-east-1"  # Specify region
        )
    )
    print(f"Index '{index_name}' created successfully!")
else:
    print(f"Index '{index_name}' already exists.")

# Fetch the host for the index
index_host = pc.describe_index(index_name).host

# Connect to the Pinecone index
pinecone_index = Index(
    name=index_name, 
    api_key=PINECONE_API_KEY,
    host=index_host
)

# Store text chunks with embeddings
docsearch = LangchainPinecone.from_documents(
    documents=text_chunks,
    embedding=embeddings,
    index_name=index_name
)

print("Data has been indexed successfully!")
